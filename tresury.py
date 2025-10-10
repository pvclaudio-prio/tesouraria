import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
from io import BytesIO
from pathlib import Path
import plotly.express as px
import os
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import zipfile
import tempfile
import json
from oauth2client.client import OAuth2Credentials
import httplib2
import traceback
import httpx
from sentence_transformers import SentenceTransformer, util
from openai import OpenAI
import requests
from difflib import SequenceMatcher
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from pandas import Timestamp
from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account
import openai
import docx
import uuid
import openpyxl
import time
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import io
import math

# =========================================
# CONFIG STREAMLIT
# =========================================
st.set_page_config(
    page_title="Revisão de Contratos - PRIO",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# =========================================
# LOGIN (mantido)
# =========================================
@st.cache_data
def carregar_usuarios():
    usuarios_config = st.secrets.get("users", {})
    usuarios = {}
    for user, dados in usuarios_config.items():
        try:
            nome, senha = dados.split("|", 1)
            usuarios[user] = {"name": nome, "password": senha}
        except:
            st.warning(f"Erro ao carregar usuário '{user}' nos secrets.")
    return usuarios

users = carregar_usuarios()

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = ""

if not st.session_state.logged_in:
    st.title("🔐 Login")
    username = st.text_input("Usuário")
    password = st.text_input("Senha", type="password")

    if st.button("Entrar"):
        user = users.get(username)
        if user and user["password"] == password:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success(f"Bem-vindo, {user['name']}!")
            st.rerun()
        else:
            st.error("Usuário ou senha incorretos.")
    st.stop()

st.sidebar.image("PRIO_SEM_POLVO_PRIO_PANTONE_LOGOTIPO_Azul.png")
nome_usuario = users[st.session_state.username]["name"]
st.sidebar.success(f"Logado como: {nome_usuario}")

if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# =========================================
# MENU
# =========================================
st.sidebar.title("Navegação")
pagina = st.sidebar.radio("Ir para:", [
    "📂 Upload do Contrato",
    "🧾 Validação das Cláusulas",
    "🔍 Análise Automática",
    "🧑‍⚖️ Revisão Final",
    "📊 Índices PRIO",
    "📘 Relatórios Gerenciais"
])

# =========================================
# GOOGLE DRIVE HELPERS
# =========================================
def conectar_drive():
    cred_dict = st.secrets["credentials"]
    credentials = OAuth2Credentials(
        access_token=cred_dict["access_token"],
        client_id=cred_dict["client_id"],
        client_secret=cred_dict["client_secret"],
        refresh_token=cred_dict["refresh_token"],
        token_expiry=datetime.strptime(cred_dict["token_expiry"], "%Y-%m-%dT%H:%M:%SZ"),
        token_uri=cred_dict["token_uri"],
        user_agent="streamlit-app/1.0",
        revoke_uri=cred_dict["revoke_uri"]
    )
    if credentials.access_token_expired:
        credentials.refresh(httplib2.Http())
    gauth = GoogleAuth()
    gauth.credentials = credentials
    return GoogleDrive(gauth)

def obter_id_pasta(nome_pasta, parent_id=None):
    drive = conectar_drive()
    query = f"title = '{nome_pasta}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    resultado = drive.ListFile({'q': query}).GetList()
    if resultado:
        return resultado[0]['id']
    return None

# =========================================
# BASES PRINCIPAIS
# - base_contratos.xlsx        (metadados do contrato)
# - clausulas_mapeadas.xlsx    (1 linha = 1 cláusula)  << NOVA BASE ROBUSTA
# - clausulas_analisadas.xlsx  (saída da análise automática)
# - clausulas_validadas.xlsx   (revisão final do usuário)
# - empresa_referencia_PRIO.xlsx (índices financeiros)
# =========================================
def carregar_base_contratos():
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    if not pasta_bases_id:
        st.error("Pasta 'bases' não encontrada.")
        return pd.DataFrame()

    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = 'base_contratos.xlsx' and trashed = false"
    }).GetList()

    if not arquivos:
        # estrutura com campos de contagem/controle
        return pd.DataFrame(columns=[
            "id_contrato", "nome_arquivo", "data_upload", "usuario_upload",
            "instituicao_financeira", "tipo", "idioma", "user_email",
            "clausulas_count"
        ])

    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    arquivos[0].GetContentFile(caminho_temp)
    df = pd.read_excel(caminho_temp)

    # garante colunas novas
    if "clausulas_count" not in df.columns:
        df["clausulas_count"] = pd.NA
    return df

def salvar_base_contratos(df):
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    pasta_backups_id = obter_id_pasta("backups", parent_id=obter_id_pasta("Tesouraria"))
    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    df.to_excel(caminho_temp, index=False)

    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = 'base_contratos.xlsx' and trashed = false"
    }).GetList()

    if arquivos:
        arquivo = arquivos[0]
    else:
        arquivo = drive.CreateFile({
            'title': 'base_contratos.xlsx',
            'parents': [{'id': pasta_bases_id}]
        })
    arquivo.SetContentFile(caminho_temp)
    arquivo.Upload()

    # Backup
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = drive.CreateFile({
        'title': f'base_contratos__{timestamp}.xlsx',
        'parents': [{'id': pasta_backups_id}]
    })
    backup.SetContentFile(caminho_temp)
    backup.Upload()

# -------- NOVAS FUNÇÕES PARA A BASE ROBUSTA DE CLÁUSULAS --------
def carregar_clausulas_mapeadas():
    """Lê clausulas_mapeadas.xlsx (1 linha = 1 cláusula)."""
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = 'clausulas_mapeadas.xlsx' and trashed = false"
    }).GetList()

    if not arquivos:
        return pd.DataFrame(columns=[
            "id_contrato", "nome_arquivo", "numero_clausula",
            "clausula", "idioma", "instituicao_financeira",
            "extraido_em", "usuario_extracao"
        ])

    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    arquivos[0].GetContentFile(caminho_temp)
    df = pd.read_excel(caminho_temp)

    # garante colunas
    for c in ["id_contrato","nome_arquivo","numero_clausula","clausula","idioma",
              "instituicao_financeira","extraido_em","usuario_extracao"]:
        if c not in df.columns:
            df[c] = ""
    return df

def salvar_clausulas_mapeadas_replace(id_contrato, df_novas_linhas):
    """
    Substitui TODAS as cláusulas do contrato (id_contrato) em clausulas_mapeadas.xlsx
    por df_novas_linhas (mesmo schema).
    """
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    pasta_backups_id = obter_id_pasta("backups", parent_id=obter_id_pasta("Tesouraria"))

    nome_arquivo = "clausulas_mapeadas.xlsx"
    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = '{nome_arquivo}' and trashed = false"
    }).GetList()

    if arquivos:
        caminho_antigo = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        arquivos[0].GetContentFile(caminho_antigo)
        df_existente = pd.read_excel(caminho_antigo)
        df_existente = df_existente[df_existente["id_contrato"] != id_contrato]
        df_final = pd.concat([df_existente, df_novas_linhas], ignore_index=True)
    else:
        df_final = df_novas_linhas

    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    df_final.to_excel(caminho_temp, index=False)

    if arquivos:
        arquivo = arquivos[0]
    else:
        arquivo = drive.CreateFile({'title': nome_arquivo, 'parents': [{'id': pasta_bases_id}]})
    arquivo.SetContentFile(caminho_temp)
    arquivo.Upload()

    # backup
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = drive.CreateFile({
        'title': f'clausulas_mapeadas__{timestamp}.xlsx',
        'parents': [{'id': pasta_backups_id}]
    })
    backup.SetContentFile(caminho_temp)
    backup.Upload()

# =========================================
# UPLOAD DE CONTRATO (mantido, sem salvar cláusulas aqui)
# =========================================
def aba_upload_contrato(user_email):
    st.title("📂 Upload do Contrato")
    st.markdown("Faça upload de um contrato em `.pdf` e preencha os dados abaixo.")

    arquivo = st.file_uploader("Selecione o contrato", type=["pdf"])
    nome_amigavel = st.text_input("Nome do contrato para exibição futura (ex: FAB PRIO - Empréstimo 2025)")
    instituicao = st.text_input("Instituição Financeira")
    idioma = st.selectbox("Idioma do Contrato", ["pt", "en"])

    if st.button("📤 Enviar Contrato"):
        if not arquivo or not nome_amigavel or not instituicao:
            st.warning("Por favor, preencha todos os campos e envie um arquivo.")
            return

        drive = conectar_drive()
        pasta_contratos_id = obter_id_pasta("contratos", parent_id=obter_id_pasta("Tesouraria"))

        id_contrato = str(uuid.uuid4())
        nome_arquivo_drive = f"{id_contrato}_{arquivo.name}"

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{arquivo.name.split('.')[-1]}") as tmp:
            tmp.write(arquivo.read())
            caminho_local = tmp.name

        novo_arquivo = drive.CreateFile({
            'title': nome_arquivo_drive,
            'parents': [{'id': pasta_contratos_id}]
        })
        novo_arquivo.SetContentFile(caminho_local)
        novo_arquivo.Upload()

        df = carregar_base_contratos()
        novo = {
            "id_contrato": id_contrato,
            "nome_arquivo": nome_amigavel,
            "tipo": arquivo.name.split(".")[-1],
            "idioma": idioma,
            "instituicao_financeira": instituicao,
            "data_upload": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "usuario_upload": user_email,
            "user_email": user_email,
            "clausulas_count": 0
        }
        df = pd.concat([df, pd.DataFrame([novo])], ignore_index=True)
        salvar_base_contratos(df)

        st.success("✅ Contrato enviado e registrado com sucesso.")

# =========================================
# MANIPULAÇÃO/EXTRAÇÃO
# =========================================
def obter_contratos_disponiveis():
    drive = conectar_drive()
    pasta_id = obter_id_pasta("contratos", parent_id=obter_id_pasta("Tesouraria"))
    arquivos = drive.ListFile({'q': f"'{pasta_id}' in parents and trashed = false"}).GetList()
    return [(arq['title'], arq['id']) for arq in arquivos]

def docx_para_pdf_temporario(caminho_docx):
    caminho_temp_dir = tempfile.mkdtemp()
    caminho_pdf = os.path.join(caminho_temp_dir, "convertido.pdf")
    convert(caminho_docx, caminho_pdf)
    return caminho_pdf

def extrair_com_document_ai_paginas(caminho_pdf, max_paginas=15):
    credentials = service_account.Credentials.from_service_account_info(st.secrets["gcp_docai"])
    project_id = st.secrets["gcp_docai"]["project_id"]
    processor_id = st.secrets["gcp_docai"]["processor_id"]
    location = "us"

    client = documentai.DocumentProcessorServiceClient(credentials=credentials)
    name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"

    leitor = PdfReader(caminho_pdf)
    total_paginas = len(leitor.pages)
    texto_total = ""

    for i in range(0, total_paginas, max_paginas):
        escritor = PdfWriter()
        for j in range(i, min(i + max_paginas, total_paginas)):
            escritor.add_page(leitor.pages[j])

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            escritor.write(temp_pdf)
            temp_pdf.flush()
            with open(temp_pdf.name, "rb") as f:
                document = {"content": f.read(), "mime_type": "application/pdf"}

        request = {"name": name, "raw_document": document}
        result = client.process_document(request=request)
        texto_total += result.document.text + "\n"

    return texto_total.strip()

def carregar_texto_contrato_drive(titulo_arquivo, arquivo_id):
    drive = conectar_drive()
    caminho_temp = tempfile.NamedTemporaryFile(delete=False).name
    drive.CreateFile({'id': arquivo_id}).GetContentFile(caminho_temp)

    try:
        if titulo_arquivo.lower().endswith(".docx"):
            caminho_pdf = docx_para_pdf_temporario(caminho_temp)
            texto = extrair_com_document_ai_paginas(caminho_pdf)
            os.remove(caminho_pdf)
        elif titulo_arquivo.lower().endswith(".pdf"):
            texto = extrair_com_document_ai_paginas(caminho_temp)
        else:
            st.error("❌ Formato de arquivo não suportado. Use .docx ou .pdf.")
            return ""
    except Exception as e:
        st.error(f"❌ Erro ao extrair o contrato: {e}")
        return ""
    return texto

# =========================================
# VALIDAÇÃO DE CLÁUSULAS (EXTRAÇÃO + SALVAR)
# =========================================
def aba_validacao_clausulas():
    st.title("🧾 Validação das Cláusulas Contratuais")

    contratos = obter_contratos_disponiveis()
    if not contratos:
        st.warning("Nenhum contrato disponível.")
        return

    nomes_arquivos = [titulo for titulo, _ in contratos]
    contrato_selecionado = st.selectbox("Selecione o contrato para análise:", nomes_arquivos)
    if not contrato_selecionado:
        st.stop()

    titulo_arquivo, id_arquivo = next(item for item in contratos if item[0] == contrato_selecionado)
    id_contrato = titulo_arquivo.split("_")[0]

    if st.session_state.get("contrato_validacao") != id_contrato:
        st.session_state["contrato_validacao"] = id_contrato
        st.session_state.pop("texto_contrato", None)
        st.session_state.pop("df_clausulas_extraidas", None)

    if st.button("▶️ Iniciar leitura do contrato"):
        with st.spinner("Lendo e extraindo texto do contrato..."):
            texto = carregar_texto_contrato_drive(titulo_arquivo, id_arquivo)
            st.session_state["texto_contrato"] = texto or ""
        if st.session_state.get("texto_contrato"):
            st.success("✅ Texto do contrato carregado com sucesso.")
        else:
            st.error("❌ Não foi possível carregar o texto do contrato.")

    if "texto_contrato" in st.session_state and st.session_state["texto_contrato"]:
        st.markdown("### 📄 Visualização do conteúdo do contrato")
        with st.expander("Visualizar texto completo extraído do contrato"):
            st.text_area("Conteúdo extraído", st.session_state["texto_contrato"], height=400)

        st.markdown("### 🧠 Passo 2 — Extrair cláusulas com IA")
        if st.button("✅ Extrair Cláusulas com IA"):
            df_clausulas = extrair_clausulas_robusto(st.session_state["texto_contrato"])
            st.session_state["df_clausulas_extraidas"] = df_clausulas

            if not df_clausulas.empty:
                st.success(f"✅ Cláusulas extraídas com sucesso! Total: {len(df_clausulas)}")
            else:
                st.warning("⚠️ Nenhuma cláusula foi extraída. Revise o texto do contrato.")
    else:
        st.info("Clique em **‘▶️ Iniciar leitura do contrato’** para carregar o texto antes de extrair as cláusulas.")

    if "df_clausulas_extraidas" in st.session_state and st.session_state["df_clausulas_extraidas"] is not None:
        st.markdown("### ✍️ Revisar Cláusulas Extraídas")
        df_editado = st.data_editor(
            st.session_state["df_clausulas_extraidas"],
            num_rows="dynamic",
            use_container_width=True,
            key="editor_clausulas"
        )

        if st.button("✅ Validar cláusulas e salvar"):
            sucesso = salvar_clausulas_validadas(df_editado, id_contrato)
            if sucesso:
                st.success("📦 Cláusulas validadas e salvas com sucesso (1 linha = 1 cláusula).")
            else:
                st.error("❌ Contrato não encontrado na base para atualização.")

# --------- CHUNKING / PROMPT / EXTRAÇÃO / DEDUPE (com ajustes) ----------
def dividir_em_chunks_simples(texto, max_chars=7000):
    if not texto:
        return []
    t = re.sub(r'\r\n?', '\n', texto)

    section_break = re.compile(
        r'(?:\n(?=[A-Z][A-Z \-\d\.\(\)]{3,}\n)|\n(?=\d+(?:\.\d+){0,3}\s)|\n(?=SECTION\s+\d+)|\n(?=DEFINITIONS)|\n(?=BACKGROUND:))',
        flags=re.IGNORECASE
    )
    parts = re.split(section_break, t)
    parts = [p.strip() for p in parts if p and p.strip()]

    chunks = []
    atual = ""
    overlap = 400  # reduzido para diminuir duplicatas

    def safe_append(acc, nxt):
        if acc:
            return acc + "\n\n" + nxt
        return nxt

    for p in parts:
        if len(atual) + len(p) + 2 <= max_chars:
            atual = safe_append(atual, p)
        else:
            if len(atual) > 0:
                chunks.append(atual.strip())
            if chunks:
                cauda = chunks[-1][-overlap:]
                atual = (cauda + "\n\n" + p).strip()
                if len(atual) > max_chars:
                    reduzir = len(atual) - max_chars + 200
                    atual = (chunks[-1][-max(0, overlap - reduzir):] + "\n\n" + p).strip()
            else:
                atual = p
            while len(atual) > max_chars:
                corte = _find_last_safe_boundary(atual, max_chars)
                chunks.append(atual[:corte].strip())
                atual = atual[corte:].strip()

    if atual:
        chunks.append(atual.strip())

    uniq = []
    seen = set()
    for c in chunks:
        key = re.sub(r'\s+', ' ', c).strip().lower()
        if key and key not in seen:
            seen.add(key)
            uniq.append(c)
    return uniq

def _find_last_safe_boundary(texto, limit):
    candidates = [
        texto.rfind("\n\n", 0, limit),
        texto.rfind(". ", 0, limit),
        texto.rfind("; ", 0, limit),
        texto.rfind("\n", 0, limit),
    ]
    pos = max([c for c in candidates if c != -1] or [limit])
    return max(1, pos)

def gerar_prompt_com_exemplos(texto_chunk):
    prompt = f"""
Você é um advogado especialista em contratos de captação de dívida (export prepayment, ECA, trade finance).
Tarefa: IDENTIFICAR e CATALOGAR **cláusulas completas** no trecho abaixo.

Regras de ouro (siga à risca):
1) **NÃO** copie exemplos, títulos do documento, marcadores, headers/footers, números de página, placeholders ([•]) ou trechos deste próprio prompt.
2) **NÃO** inclua numeração/títulos. Extraia **apenas o texto integral da cláusula**.
3) **NÃO** resuma. **NÃO** reescreva. **NÃO** traduza. Retorne o texto **exato** da cláusula conforme o contrato.
4) Considere como "cláusula" todo enunciado normativo/operacional **completo** (obrigações, definições, prazos, taxas, eventos de default, lei aplicável, etc.) que possa ser referenciado isoladamente.
   - Em definições, capture o enunciado inteiro.
   - Em listas (a), (b), (c) que formam uma cláusula única, una os itens em **uma única string**.
5) **NÃO** quebre cláusulas em várias saídas; **cada item do array deve conter uma cláusula completa**.
6) Se o trecho contém apenas parte de uma cláusula (foi cortada), **ignore** essa cláusula neste chunk.
7) Saída **apenas** em JSON válido, no formato:
{{
  "clauses": [
    "cláusula 1 (texto completo sem título/numeração)",
    "cláusula 2"
  ]
}}

Agora processe o trecho a seguir:

\"\"\"{texto_chunk}\"\"\"    
"""
    return prompt.strip()

def extrair_clausulas_robusto(texto):
    client = OpenAI(api_key=st.secrets["openai"]["api_key"])
    st.info("🔍 Analisando o contrato...")
    partes = dividir_em_chunks_simples(texto)
    clausulas_total = []
    total_ok = 0

    for i, chunk in enumerate(partes):
        with st.spinner(f"Extraindo cláusulas do contrato: {i+1}/{len(partes)}..."):
            prompt = gerar_prompt_com_exemplos(chunk)
            try:
                resposta = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": (
                            "Você é um especialista jurídico com ampla experiência em contratos "
                            "de dívida e operações de pré-pagamento de exportação. Siga estritamente as instruções do usuário."
                        )},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0,
                    max_tokens=4096
                )
                saida = (resposta.choices[0].message.content or "").strip()
                clausulas = _parse_clauses_from_output(saida)
                total_ok += len(clausulas)
                clausulas_total.extend(clausulas)
                st.write(f"Chunk {i+1}: {len(clausulas)} cláusulas válidas")
            except Exception as e:
                st.error(f"Erro no chunk {i+1}: {e}")

    clausulas_total = _dedupe_clauses(clausulas_total, sim_threshold=0.92)
    st.info(f"🔎 Total após deduplicação: {len(clausulas_total)} cláusulas")

    df = pd.DataFrame(clausulas_total, columns=["clausula"])
    # adiciona numeração sequencial para inspeção na UI
    df.insert(0, "numero_clausula", range(1, len(df) + 1))
    return df

def _parse_clauses_from_output(saida: str):
    try:
        data = json.loads(saida)
        if isinstance(data, dict) and "clauses" in data and isinstance(data["clauses"], list):
            return [_clean_clause_text(c) for c in data["clauses"] if isinstance(c, str) and _clean_clause_text(c)]
    except Exception:
        pass

    try:
        match = re.search(r'\{[\s\S]*\}', saida)
        if match:
            data = json.loads(match.group(0))
            if isinstance(data, dict) and "clauses" in data and isinstance(data["clauses"], list):
                return [_clean_clause_text(c) for c in data["clauses"] if isinstance(c, str) and _clean_clause_text(c)]
    except Exception:
        pass

    linhas = [l.strip() for l in re.split(r'\n{2,}', saida) if l.strip()]
    linhas = [l for l in linhas if len(l) > 30]
    return [_clean_clause_text(l) for l in linhas if _clean_clause_text(l)]

def _clean_clause_text(txt: str) -> str:
    t = txt.strip()
    t = re.sub(r'^[0-9]+(\.[0-9]+)*\s*[-–—]*\s*', '', t)
    t = re.sub(r'^(SECTION|SEÇÃO|ARTIGO|CLAUSE)\s+[0-9A-Za-z\.\-–—]+\s*[:\-–—]\s*', '', t, flags=re.I)
    t = re.sub(r'\b\d{5,}v\d+\b', '', t)  # ex: 13348400v3
    t = re.sub(r'\[[^\]]*\]', '', t)      # [•]
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\s+\n', '\n', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

def _norm_for_hash(s: str) -> str:
    s2 = s.lower()
    s2 = re.sub(r'\s+', ' ', s2)
    s2 = re.sub(r'["“”\'`´]', '', s2)
    s2 = re.sub(r'[\(\)\[\]\{\}]', '', s2)
    s2 = re.sub(r'\bsection\b|\bclause\b|\bartigo\b', '', s2)
    s2 = re.sub(r'\d{5,}v\d+', '', s2)
    return s2.strip()

def _similar(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()

def _dedupe_clauses(clausulas, sim_threshold=0.92):
    uniq = []
    for c in clausulas:
        cc = _clean_clause_text(c)
        if not cc:
            continue
        n = _norm_for_hash(cc)
        dup_idx = -1
        best_sim = 0.0
        for i, u in enumerate(uniq):
            sim = _similar(n, _norm_for_hash(u))
            if sim > best_sim:
                best_sim = sim
                dup_idx = i
        if best_sim >= sim_threshold and abs(len(cc) - len(uniq[dup_idx])) < max(120, 0.15*len(uniq[dup_idx])):
            if len(cc) > len(uniq[dup_idx]):
                uniq[dup_idx] = cc
        else:
            uniq.append(cc)
    return uniq

# --------- SALVAMENTO ROBUSTO (1 linha = 1 cláusula) ----------
def salvar_clausulas_validadas(df_clausulas, id_contrato):
    """
    Agora salva em clausulas_mapeadas.xlsx (uma linha por cláusula)
    e atualiza a contagem em base_contratos.xlsx.
    """
    base = carregar_base_contratos()
    if base.empty:
        return False
    idx = base[base["id_contrato"] == id_contrato].index
    if len(idx) == 0:
        return False

    nome_arquivo = base.loc[idx[0], "nome_arquivo"]
    idioma = base.loc[idx[0], "idioma"] if "idioma" in base.columns else ""
    instituicao = base.loc[idx[0], "instituicao_financeira"] if "instituicao_financeira" in base.columns else ""
    usuario = base.loc[idx[0], "usuario_upload"] if "usuario_upload" in base.columns else st.session_state.username

    df_clausulas = df_clausulas.copy()
    if "clausula" not in df_clausulas.columns:
        st.error("DataFrame de cláusulas não possui coluna 'clausula'.")
        return False

    df_clausulas["clausula"] = df_clausulas["clausula"].astype(str)
    if "numero_clausula" not in df_clausulas.columns:
        df_clausulas.insert(0, "numero_clausula", range(1, len(df_clausulas) + 1))

    df_novas = pd.DataFrame({
        "id_contrato": id_contrato,
        "nome_arquivo": nome_arquivo,
        "numero_clausula": df_clausulas["numero_clausula"].astype(int),
        "clausula": df_clausulas["clausula"].astype(str),
        "idioma": idioma,
        "instituicao_financeira": instituicao,
        "extraido_em": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "usuario_extracao": usuario
    })

    salvar_clausulas_mapeadas_replace(id_contrato, df_novas)

    # Atualiza contagem em base_contratos
    base.loc[idx[0], "clausulas_count"] = len(df_novas)
    salvar_base_contratos(base)
    return True

# =========================================
# CARREGAR CLÁUSULAS PARA ANÁLISE (agora desde clausulas_mapeadas.xlsx)
# =========================================
def carregar_clausulas_contratos():
    df_map = carregar_clausulas_mapeadas()
    if df_map.empty:
        return pd.DataFrame(columns=["nome_arquivo", "clausula"])
    return df_map[["nome_arquivo", "clausula"]].copy()

# =========================================
# ANÁLISE AUTOMÁTICA
# =========================================
def carregar_clausulas_analisadas():
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))

    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = 'clausulas_analisadas.xlsx' and trashed = false"
    }).GetList()

    if not arquivos:
        st.warning("❌ Base de cláusulas analisadas não encontrada.")
        return pd.DataFrame(columns=[
            "id_contrato","nome_arquivo","clausula",
            "revisao_juridico","motivo_juridico",
            "revisao_financeiro","motivo_financeiro",
            "revisao_sup","motivo_sup",
            "run_id","analisado_em"
        ])

    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    arquivos[0].GetContentFile(caminho_temp)
    df = pd.read_excel(caminho_temp)

    rename_map = {
        "analise_juridico_status": "revisao_juridico",
        "analise_juridico_motivo": "motivo_juridico",
        "analise_financeiro_status": "revisao_financeiro",
        "analise_financeiro_motivo": "motivo_financeiro",
        "revisao_juridico_status": "revisao_juridico",
        "revisao_juridico_motivo": "motivo_juridico",
        "revisao_financeiro_status": "revisao_financeiro",
        "revisao_financeiro_motivo": "motivo_financeiro",
        "revisao_sup_status": "revisao_sup",
        "revisao_sup_motivo": "motivo_sup",
    }
    df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})

    for c in ["id_contrato","revisao_juridico","motivo_juridico","revisao_financeiro","motivo_financeiro","revisao_sup","motivo_sup","run_id","analisado_em"]:
        if c not in df.columns:
            df[c] = ""
    return df

def salvar_clausulas_validadas_usuario(df_novo):
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    pasta_backups_id = obter_id_pasta("backups", parent_id=obter_id_pasta("Tesouraria"))

    nome_arquivo = "clausulas_analisadas.xlsx"
    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name

    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = '{nome_arquivo}' and trashed = false"
    }).GetList()

    if arquivos:
        caminho_antigo = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        arquivos[0].GetContentFile(caminho_antigo)
        df_existente = pd.read_excel(caminho_antigo)

        contrato_atual = df_novo["nome_arquivo"].iloc[0]
        df_existente = df_existente[df_existente["nome_arquivo"] != contrato_atual]

        df_final = pd.concat([df_existente, df_novo], ignore_index=True)
    else:
        df_final = df_novo

    df_final.to_excel(caminho_temp, index=False)

    if arquivos:
        arquivo = arquivos[0]
    else:
        arquivo = drive.CreateFile({
            'title': nome_arquivo,
            'parents': [{'id': pasta_bases_id}]
        })
    arquivo.SetContentFile(caminho_temp)
    arquivo.Upload()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = drive.CreateFile({
        'title': f'clausulas_analisadas__{timestamp}.xlsx',
        'parents': [{'id': pasta_backups_id}]
    })
    backup.SetContentFile(caminho_temp)
    backup.Upload()

def aba_analise_automatica():
    st.title("🧠 Análise Automática das Cláusulas")

    df = carregar_clausulas_contratos()  # agora vem de clausulas_mapeadas.xlsx
    df_contrato = carregar_clausulas_analisadas()

    contratos_disponiveis = df["nome_arquivo"].dropna().unique().tolist()
    contrato_escolhido = st.selectbox("Selecione o contrato:", contratos_disponiveis)

    df_clausulas = df[df["nome_arquivo"] == contrato_escolhido].copy() if contrato_escolhido else pd.DataFrame()
    clausulas = [c.strip() for c in df_clausulas["clausula"].tolist() if c.strip()] if not df_clausulas.empty else []

    if clausulas:
        if st.button("✅ Iniciar Análise Automática"):
            drive = conectar_drive()
            pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
            arquivos = drive.ListFile({
                'q': f"'{pasta_bases_id}' in parents and title = 'empresa_referencia_PRIO.xlsx' and trashed = false"
            }).GetList()
            if not arquivos:
                st.error("Base de índices financeiros 'empresa_referencia_PRIO.xlsx' não encontrada.")
                return

            caminho_indices = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
            arquivos[0].GetContentFile(caminho_indices)
            df_indices = pd.read_excel(caminho_indices)

            client = OpenAI(api_key=st.secrets["openai"]["api_key"])
            resultados = []
            st.info("🔍 Iniciando análise com os especialistas jurídico e financeiro...")

            progress_bar = st.progress(0)
            status_text = st.empty()

            # recuperar id_contrato a partir do nome no diretório de contratos
            contratos_drive = obter_contratos_disponiveis()
            id_contrato_sel = next((x.split("_")[0] for x, _ in contratos_drive if x == contrato_escolhido), "")

            run_id = datetime.now().strftime("%Y%m%d_%H%M%S")

            for i, clausula in enumerate(clausulas):
                status_text.text(f"Processando cláusula {i+1}/{len(clausulas)}...")
                with st.spinner():
                    prompt_juridico = f"""
Você é um advogado especialista em contratos de dívida.
Analise a cláusula abaixo e diga se está Conforme ou se Necessita Revisão. Você somente pode escolher uma alternativa.
Sempre inicie sua resposta com exatamente as palavras Conforme ou Necessita Revisão.
Justifique de forma objetiva com base jurídica.

Cláusula:
\"\"\"{clausula}\"\"\""""
                    resposta_juridico = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": prompt_juridico}],
                        temperature=0,
                        max_tokens=1000
                    ).choices[0].message.content.strip()

                    texto_indices = df_indices.to_string(index=False)
                    prompt_financeiro = f"""
Você é um especialista financeiro com foco em contratos de captação de dívida. Abaixo estão os índices financeiros da empresa PRIO:

{texto_indices}

Analise a cláusula a seguir e diga se ela está financeiramente Conforme ou se Necessita Revisão. Você somente pode escolher uma alternativa.
Sempre inicie sua resposta com exatamente as palavras Conforme ou Necessita Revisão.
Caso a cláusula não aborde nenhuma condicionante financeira, diga que está Conforme e no motivo informe objetivamente que não foram identificados
índices financeiros para análise.
Justifique com base nos dados da empresa e benchmarking de mercado.

Cláusula:
\"\"\"{clausula}\"\"\""""
                    resposta_financeiro = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": prompt_financeiro}],
                        temperature=0,
                        max_tokens=1000
                    ).choices[0].message.content.strip()

                    prompt_supervisor = f"""
Você é o supervisor responsável pela revisão final. 
Abaixo está a cláusula, a análise do agente jurídico e a análise do agente financeiro. 
Revise cada uma delas e diga se Concorda ou Não Concorda, e explique brevemente o motivo.
Sempre inicie sua resposta com exatamente as palavras Concorda ou Não Concorda.

Cláusula:
\"\"\"{clausula}\"\"\"


Análise Jurídica:
{resposta_juridico}

Análise Financeira:
{resposta_financeiro}"""
                    resposta_supervisor = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": prompt_supervisor}],
                        temperature=0,
                        max_tokens=1000
                    ).choices[0].message.content.strip()

                    jur_raw = (resposta_juridico or "").strip().lower()
                    jur_status = "Conforme" if jur_raw.startswith("conforme") else "Necessita Revisão"

                    fin_raw = (resposta_financeiro or "").strip().lower()
                    fin_status = "Conforme" if fin_raw.startswith("conforme") else "Necessita Revisão"

                    sup_raw = (resposta_supervisor or "").strip().lower()
                    if sup_raw.startswith("não concorda"):
                        sup_status = "Não Concorda"
                    elif sup_raw.startswith("concorda"):
                        sup_status = "Concorda"
                    else:
                        sup_status = "Não Concorda"

                    resultados.append({
                        "id_contrato": id_contrato_sel,
                        "nome_arquivo": contrato_escolhido,
                        "clausula": clausula,
                        "revisao_juridico": jur_status,
                        "motivo_juridico": resposta_juridico,
                        "revisao_financeiro": fin_status,
                        "motivo_financeiro": resposta_financeiro,
                        "revisao_sup": sup_status,
                        "motivo_sup": resposta_supervisor,
                        "run_id": run_id,
                        "analisado_em": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    })

                progress_bar.progress((i + 1) / len(clausulas))

            df_resultado = pd.DataFrame(resultados)
            st.session_state["analise_automatica_resultado"] = df_resultado
            st.success("✅ Análise automática concluída.")

    else:
        st.warning("Não há cláusulas validadas disponíveis.")

    if "analise_automatica_resultado" in st.session_state:
        df_resultado = st.session_state["analise_automatica_resultado"]
        st.dataframe(df_resultado, use_container_width=True)

        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df_resultado.to_excel(writer, index=False)

        st.download_button("📥 Baixar Análises", data=buffer.getvalue(),
                           file_name="clausulas_analisadas.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key="download_novo")

        if st.button("Desejar Salvar ?", key="salvar_novo"):
            salvar_clausulas_validadas_usuario(df_resultado)
            st.success("✅ Revisão dos agentes foi salva com sucesso!")
            del st.session_state["analise_automatica_resultado"]

    elif df_contrato is not None and not df_contrato.empty and contrato_escolhido:
        df_contrato = df_contrato[df_contrato["nome_arquivo"] == contrato_escolhido]
        if not df_contrato.empty:
            st.dataframe(df_contrato, use_container_width=True)
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                df_contrato.to_excel(writer, index=False)
            st.download_button("📥 Baixar Análises", data=buffer.getvalue(),
                               file_name="clausulas_analisadas.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               key="download_anterior")

# =========================================
# REVISÃO FINAL (mantido c/ pequenas garantias)
# =========================================
USER_REVISAO_OPCOES = ["Concordo", "Discordo", "Melhoria"]

def aba_revisao_final():
    st.title("🧑‍⚖️ Revisão Final do Usuário - Cláusulas Contratuais")

    st.markdown("""
        <style>
        .stDataFrame td, .stDataFrame div, .stDataEditor td, .stDataEditor div {
            white-space: normal !important;
        }
        .stDataEditor [data-testid="stVerticalBlock"] { overflow: visible !important; }
        </style>
    """, unsafe_allow_html=True)

    with st.spinner("Carregando cláusulas analisadas..."):
        df = carregar_clausulas_analisadas()
    if df is None or df.empty:
        st.warning("Nenhuma cláusula analisada disponível.")
        return

    contratos_disponiveis = df["nome_arquivo"].dropna().unique().tolist()
    contrato = st.selectbox("Selecione o contrato:", contratos_disponiveis)
    if not contrato:
        return

    df_filtrado = df[df["nome_arquivo"] == contrato].copy()
    # garante presença e tipo de colunas editáveis
    for col in ["user_revisao", "motivo_user"]:
        if col not in df_filtrado.columns:
            df_filtrado[col] = ""
    
    # remove objetos/None/NaN que quebram o React (tudo vira string segura)
    safe_cols = [
        "clausula","revisao_juridico","motivo_juridico",
        "revisao_financeiro","motivo_financeiro",
        "revisao_sup","motivo_sup",
        "user_revisao","motivo_user",
        "id_contrato","nome_arquivo","run_id","analisado_em"
    ]
    for c in safe_cols:
        if c in df_filtrado.columns:
            df_filtrado[c] = df_filtrado[c].astype(str).fillna("")
    
    # evita valores fora das opções do SelectboxColumn
    if "user_revisao" in df_filtrado.columns:
        valid_opts = set(["Concordo","Discordo","Melhoria",""])
        df_filtrado["user_revisao"] = df_filtrado["user_revisao"].apply(
            lambda x: x if x in valid_opts else ""
        )

    st.markdown("### 📝 Revisão Final do Usuário")

    for col in ["user_revisao", "motivo_user"]:
        if col not in df_filtrado.columns:
            df_filtrado[col] = ""

    colunas_ordem = [
        "clausula",
        "revisao_juridico", "motivo_juridico",
        "revisao_financeiro", "motivo_financeiro",
        "revisao_sup", "motivo_sup",
        "user_revisao", "motivo_user",
        "id_contrato", "nome_arquivo",
    ]
    colunas_ordem = [c for c in colunas_ordem if c in df_filtrado.columns]

    col_cfg = {
        "clausula": st.column_config.TextColumn("Cláusula", width="large"),
        "motivo_juridico": st.column_config.TextColumn("Motivo Jurídico", width="large"),
        "motivo_financeiro": st.column_config.TextColumn("Motivo Financeiro", width="large"),
        "motivo_sup": st.column_config.TextColumn("Motivo Supervisor", width="large"),
        "user_revisao": st.column_config.SelectboxColumn(
            "Revisão do Usuário",
            options=USER_REVISAO_OPCOES,
            help="Selecione sua revisão para a cláusula"
        ),
        "motivo_user": st.column_config.TextColumn(
            "Motivo (usuário)",
            help="Explique de forma objetiva sua concordância/discordância ou sugestão de melhoria",
            width="large"
        ),
        "nome_arquivo": st.column_config.TextColumn("Contrato (interno)"),
        "revisao_juridico": st.column_config.TextColumn("Revisão Jurídica"),
        "revisao_financeiro": st.column_config.TextColumn("Revisão Financeira"),
        "revisao_sup": st.column_config.TextColumn("Revisão Supervisor"),
    }

    desabilitadas = [c for c in df_filtrado.columns if c not in ["user_revisao", "motivo_user"]]

    linhas = len(df_filtrado)
    altura = min(700, 56 + 40 * min(15, linhas))

    with st.form("form_revisao_final", clear_on_submit=False):
        df_editado = st.data_editor(
            df_filtrado,
            column_config=col_cfg,
            column_order=colunas_ordem,
            disabled=desabilitadas,
            hide_index=True,
            num_rows="fixed",
            use_container_width=True,
            height=altura,
            key="revisao_final_editor"
        )

        col_a, col_b = st.columns([1, 2])
        salvar_click = col_a.form_submit_button("✅ Salvar revisão final do usuário", use_container_width=True)
        baixar_click = col_b.form_submit_button("⬇️ Baixar análises (.xlsx)", use_container_width=True)

    if baixar_click:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df_editado.to_excel(writer, index=False)
        st.download_button(
            "📥 Clique para baixar",
            data=buffer.getvalue(),
            file_name="clausulas_validadas.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

    if salvar_click:
        try:
            salvar_clausulas_revisadas_usuario(df_editado)
            st.success("✅ Revisão final do usuário salva com sucesso!")
        except Exception as e:
            st.error(f"Falha ao salvar no Drive: {e}")

    st.caption(f"A base possui **{linhas}** cláusulas para o contrato selecionado.")

def salvar_clausulas_revisadas_usuario(df_novo: pd.DataFrame):
    drive = conectar_drive()
    pasta_tesouraria_id = obter_id_pasta("Tesouraria")
    pasta_bases_id = obter_id_pasta("bases", parent_id=pasta_tesouraria_id)
    pasta_backups_id = obter_id_pasta("backups", parent_id=pasta_tesouraria_id)

    nome_arquivo = "clausulas_validadas.xlsx"
    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name

    arquivos = drive.ListFile({
        "q": f"'{pasta_bases_id}' in parents and title = '{nome_arquivo}' and trashed = false"
    }).GetList()

    if arquivos:
        caminho_antigo = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        arquivos[0].GetContentFile(caminho_antigo)
        df_existente = pd.read_excel(caminho_antigo)

        contrato_atual = df_novo["nome_arquivo"].iloc[0]
        df_existente = df_existente[df_existente["nome_arquivo"] != contrato_atual]

        df_final = pd.concat([df_existente, df_novo], ignore_index=True)
    else:
        df_final = df_novo

    df_final.to_excel(caminho_temp, index=False)

    if arquivos:
        arquivo = arquivos[0]
    else:
        arquivo = drive.CreateFile({"title": nome_arquivo, "parents": [{"id": pasta_bases_id}]})
    arquivo.SetContentFile(caminho_temp)
    arquivo.Upload()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = drive.CreateFile({
        "title": f"clausulas_validadas__{timestamp}.xlsx",
        "parents": [{"id": pasta_backups_id}]
    })
    backup.SetContentFile(caminho_temp)
    backup.Upload()

def carregar_clausulas_validadas():
    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))

    arquivos = drive.ListFile({
        "q": f"'{pasta_bases_id}' in parents and title = 'clausulas_validadas.xlsx' and trashed = false"
    }).GetList()

    if not arquivos:
        st.warning("❌ Base de cláusulas validadas não encontrada.")
        return pd.DataFrame(columns=[
            "id_contrato","nome_arquivo", "clausula",
            "revisao_juridico", "motivo_juridico",
            "revisao_financeiro", "motivo_financeiro",
            "revisao_sup", "motivo_sup",
            "user_revisao", "motivo_user"
        ])

    caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
    arquivos[0].GetContentFile(caminho_temp)
    return pd.read_excel(caminho_temp)

# =========================================
# ÍNDICES PRIO (mantido)
# =========================================
def aba_indices_prio():
    st.title("📊 Índices Financeiros da PRIO")

    drive = conectar_drive()
    pasta_bases_id = obter_id_pasta("bases", parent_id=obter_id_pasta("Tesouraria"))
    pasta_backups_id = obter_id_pasta("backups", parent_id=obter_id_pasta("Tesouraria"))

    nome_arquivo = "empresa_referencia_PRIO.xlsx"

    arquivos = drive.ListFile({
        'q': f"'{pasta_bases_id}' in parents and title = '{nome_arquivo}' and trashed = false"
    }).GetList()

    if not arquivos:
        st.warning("Base 'empresa_referencia_PRIO.xlsx' não encontrada. Será criada uma nova base.")
        df_indices = pd.DataFrame(columns=["EBITDA", "Mrg EBITDA", "Res Fin", "Dívida", "Lucro Líq", "Caixa"])
    else:
        caminho_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        arquivos[0].GetContentFile(caminho_temp)
        df_indices = pd.read_excel(caminho_temp)

    st.markdown("### ✍️ Editar Índices")
    df_editado = st.data_editor(
        df_indices,
        num_rows="dynamic",
        use_container_width=True,
        key="editor_indices_prio"
    )

    if st.button("💾 Salvar Índices"):
        caminho_temp_salvar = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx").name
        df_editado.to_excel(caminho_temp_salvar, index=False)

        if arquivos:
            arquivo = arquivos[0]
        else:
            arquivo = drive.CreateFile({
                'title': nome_arquivo,
                'parents': [{'id': pasta_bases_id}]
            })

        arquivo.SetContentFile(caminho_temp_salvar)
        arquivo.Upload()

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup = drive.CreateFile({
            'title': f"empresa_referencia_PRIO__{timestamp}.xlsx",
            'parents': [{'id': pasta_backups_id}]
        })
        backup.SetContentFile(caminho_temp_salvar)
        backup.Upload()

        st.success("✅ Índices salvos e backup criado com sucesso!")

# =========================================
# RELATÓRIO GERENCIAL (mantido)
# =========================================
def aba_relatorios_gerenciais():
    st.title("📘 Relatório Gerencial - Ações Prioritárias por Contrato")

    df = carregar_clausulas_validadas()
    if df.empty:
        st.warning("Base de cláusulas validadas está vazia.")
        return

    contratos = df["nome_arquivo"].unique().tolist()
    contrato_selecionado = st.selectbox("Selecione o contrato para análise:", contratos)

    if not contrato_selecionado:
        return

    if st.button("✅ Executar análise"):
        clausulas_contrato = df[df["nome_arquivo"] == contrato_selecionado]["clausula"].tolist()

        texto_clausulas = "\n\n".join(clausulas_contrato)
        prompt = f"""
Você é um especialista jurídico em gestão contratual e compliance.

Com base nas cláusulas abaixo, elenque de forma objetiva e por ordem de significância as principais ações que o usuário deve realizar para garantir a segurança jurídica do contrato.

As ações precisam ser específicas para as cláusulas do contrato marcadas como Necessita Revisão. Não traga ações generalistas, seja crítico, objetivo e especialista.

Sua resposta deve conter no máximo 1 página e apresentar as ações com títulos curtos, seguidos de explicações objetivas (1 parágrafo por ação). Seja direto, técnico e evite repetições.

Mantenha sempre uma breve referência à cláusula que precisa ser revisada para assegurar a conformidade.

Cláusulas do contrato:
\"\"\"{texto_clausulas}\"\"\""""

        client = OpenAI(api_key=st.secrets["openai"]["api_key"])
        with st.spinner("Gerando análise..."):
            resposta = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Você é um consultor jurídico especialista em contratos de captação de dívida internacionais."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0,
                max_tokens=2048
            )

        analise_final = resposta.choices[0].message.content.strip()
        st.markdown("### ✅ Análise Gerada:")
        st.markdown(analise_final)

        buffer = BytesIO()
        doc = Document()
        doc.add_heading(f"Relatório Gerencial - {contrato_selecionado}", level=1)
        for par in analise_final.split("\n"):
            if par.strip():
                doc.add_paragraph(par.strip())
        doc.save(buffer)
        st.download_button(
            label="📥 Baixar Análise",
            data=buffer.getvalue(),
            file_name=f"relatorio_gerencial_{contrato_selecionado}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# =========================================
# ROUTER
# =========================================
if pagina == "📂 Upload do Contrato":
    aba_upload_contrato(user_email=st.session_state.username)
elif pagina == "🧾 Validação das Cláusulas":
    aba_validacao_clausulas()
elif pagina == "🔍 Análise Automática":
    aba_analise_automatica()
elif pagina == "🧑‍⚖️ Revisão Final":
    aba_revisao_final()
elif pagina == "📊 Índices PRIO":
    aba_indices_prio()
elif pagina == "📘 Relatórios Gerenciais":
    aba_relatorios_gerenciais()
